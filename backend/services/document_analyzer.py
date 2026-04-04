from typing import List, Dict, Any, Optional
import httpx
import logging
from pathlib import Path
import json
from core.config import settings
from services.a2a_client import execute_via_a2a
from services.httpx_tls import httpx_verify_parameter
from services.job_file_storage import materialize_to_temp_path, cleanup_temp_path
from services.planner_llm import is_agent_planner_configured, planner_chat_completion

logger = logging.getLogger(__name__)


class DocumentAnalyzer:
    """Extract document text and run Q&A via hired agent."""

    def __init__(self):
        pass

    async def read_document(self, file_path: str) -> str:
        """Read document content based on file type and extract text for inference endpoint"""
        path = Path(file_path)
        file_ext = path.suffix.lower()
        
        try:
            # Text-based files
            if file_ext in ['.txt', '.md', '.rtf']:
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif file_ext == '.csv':
                import csv
                content = []
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        content.append(','.join(row))
                return '\n'.join(content)
            
            elif file_ext == '.json':
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return json.dumps(data, indent=2)
            
            elif file_ext == '.xml':
                with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            # PDF files
            elif file_ext == '.pdf':
                try:
                    from pypdf import PdfReader
                    content = []
                    with open(path, 'rb') as f:
                        pdf_reader = PdfReader(f)
                        for page_num, page in enumerate(pdf_reader.pages):
                            text = page.extract_text()
                            if text.strip():
                                content.append(f"--- Page {page_num + 1} ---\n{text}")
                    return '\n\n'.join(content) if content else "[PDF file contains no extractable text]"
                except ImportError:
                    return f"[PDF extraction requires pypdf library. File: {path.name}]"
                except Exception as e:
                    return f"[Error extracting PDF content: {str(e)}]"
            
            # Word documents (.docx) - try python-docx first, then docx2txt fallback
            elif file_ext == '.docx':
                # Primary: python-docx (better structure for paragraphs/tables)
                try:
                    from docx import Document
                    doc = Document(path)
                    content = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content.append(paragraph.text)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                content.append(row_text)
                    if content:
                        return '\n'.join(content)
                    # Empty body - try docx2txt which can get more content (headers, footers)
                except ImportError:
                    pass
                except Exception:
                    pass
                # Fallback: docx2txt (works without python-docx, extracts text from .docx)
                try:
                    import docx2txt
                    text = docx2txt.process(str(path))
                    return text.strip() if text and text.strip() else "[DOCX file contains no extractable text]"
                except ImportError:
                    return f"[DOCX extraction requires python-docx or docx2txt. File: {path.name}]"
                except Exception as e:
                    return f"[Error extracting DOCX content: {str(e)}]"
            
            # Legacy Word documents (.doc) - requires additional library
            elif file_ext == '.doc':
                try:
                    # Try using python-docx2txt or textract
                    import docx2txt
                    text = docx2txt.process(path)
                    return text if text.strip() else "[DOC file contains no extractable text]"
                except ImportError:
                    return f"[DOC extraction requires docx2txt library. File: {path.name}]"
                except Exception as e:
                    return f"[Error extracting DOC content: {str(e)}]"
            
            # Excel files (.xlsx, .xls)
            elif file_ext in ['.xlsx', '.xls']:
                try:
                    import pandas as pd
                    # Read all sheets
                    excel_file = pd.ExcelFile(path)
                    content = []
                    for sheet_name in excel_file.sheet_names:
                        df = pd.read_excel(path, sheet_name=sheet_name)
                        content.append(f"--- Sheet: {sheet_name} ---")
                        # Convert DataFrame to string representation
                        content.append(df.to_string(index=False))
                    return '\n\n'.join(content) if content else "[Excel file contains no data]"
                except ImportError:
                    return f"[Excel extraction requires pandas and openpyxl libraries. File: {path.name}]"
                except Exception as e:
                    return f"[Error extracting Excel content: {str(e)}]"
            
            # OpenDocument formats (.odt, .ods)
            elif file_ext == '.odt':
                try:
                    from odf import text, teletype
                    from odf.opendocument import load
                    doc = load(path)
                    paragraphs = doc.getElementsByType(text.P)
                    content = []
                    for para in paragraphs:
                        text_content = teletype.extractText(para)
                        if text_content.strip():
                            content.append(text_content)
                    return '\n'.join(content) if content else "[ODT file contains no extractable text]"
                except ImportError:
                    return f"[ODT extraction requires odfpy library. File: {path.name}]"
                except Exception as e:
                    return f"[Error extracting ODT content: {str(e)}]"
            
            elif file_ext == '.ods':
                try:
                    import pandas as pd
                    df = pd.read_excel(path, engine='odf')
                    return df.to_string(index=False) if not df.empty else "[ODS file contains no data]"
                except ImportError:
                    return f"[ODS extraction requires pandas and odfpy libraries. File: {path.name}]"
                except Exception as e:
                    return f"[Error extracting ODS content: {str(e)}]"
            
            else:
                return f"[Unsupported file type: {file_ext}. File: {path.name}]"
                
        except Exception as e:
            return f"[Error reading file {path.name}: {str(e)}]"

    async def read_file_info(self, file_info: Dict[str, Any]) -> str:
        """Read document content from metadata entry (local path or S3-backed object)."""
        local_path = await materialize_to_temp_path(file_info)
        try:
            return await self.read_document(local_path)
        finally:
            cleanup_temp_path(file_info, local_path)
    
    async def _analyze_documents_via_a2a(
        self,
        all_content: str,
        job_title: str,
        job_description: Optional[str],
        conversation_history: List[Dict[str, str]],
        agent_api_url: str,
        agent_api_key: Optional[str],
        *,
        adapter_url: Optional[str] = None,
        adapter_metadata: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Call hired agent via A2A for BRD analysis (direct or via platform adapter); return same shape as OpenAI path."""
        input_data = {
            "job_title": job_title,
            "job_description": job_description or "",
            "documents_content": all_content,
            "conversation_history": conversation_history,
            "task": "Analyze the documents and requirements. Ask ONLY important clarifying questions (critical gaps, real ambiguities, risks). Do NOT ask minor or obvious questions (e.g. for simple math like 'add 2+5' do not ask output format—integer/float; use a reasonable default). Return valid JSON: {\"analysis\": \"...\", \"questions\": [] (only if truly needed), \"recommendations\": [], \"solutions\": [], \"next_steps\": []}. Works for both sequential and A2A workflows.",
        }
        if adapter_url and adapter_metadata is not None:
            result = await execute_via_a2a(
                adapter_url,
                input_data,
                api_key=None,
                blocking=True,
                timeout=120.0,
                adapter_metadata=adapter_metadata,
            )
        else:
            result = await execute_via_a2a(
                agent_api_url,
                input_data,
                api_key=agent_api_key,
                blocking=True,
                timeout=120.0,
            )
        content = (result.get("content") or "").strip()
        try:
            parsed = json.loads(content)
            hint = parsed.get("workflow_collaboration_hint")
            if hint not in ("sequential", "async_a2a"):
                hint = None
            return {
                "analysis": parsed.get("analysis", ""),
                "questions": parsed.get("questions", []),
                "recommendations": parsed.get("recommendations", []),
                "solutions": parsed.get("solutions", []),
                "next_steps": parsed.get("next_steps", []),
                "workflow_collaboration_hint": hint,
                "workflow_collaboration_reason": (parsed.get("workflow_collaboration_reason") or "").strip() or None if hint else None,
                "raw_response": content,
            }
        except json.JSONDecodeError:
            return {
                "analysis": content,
                "questions": self._extract_questions(content),
                "recommendations": self._extract_recommendations(content),
                "solutions": [],
                "next_steps": [],
                "workflow_collaboration_hint": None,
                "workflow_collaboration_reason": None,
                "raw_response": content,
            }

    async def _analyze_with_platform_planner(
        self,
        *,
        all_content: str,
        job_title: str,
        job_description: Optional[str],
        conversation_history: List[Dict[str, str]],
    ) -> Dict[str, Any]:
        """BRD analysis via platform-configured planner model (Issue #62)."""
        system_prompt = """You are an expert AI assistant specialized in deeply understanding business requirements from documents and providing intelligent solutions.

Your primary tasks:
1. DEEP ANALYSIS: Thoroughly analyze all uploaded documents to extract:
   - Business objectives and goals
   - Technical requirements and specifications
   - Data structures, formats, and sources
   - Workflow processes and dependencies
   - Constraints, limitations, and edge cases
   - Success criteria and expected outcomes

2. INTELLIGENT QUESTIONING: Ask ONLY important, high-impact questions that:
   - Fill critical gaps in understanding (missing scope, constraints, or success criteria)
   - Clarify genuinely ambiguous requirements that affect the outcome
   - Identify real risks or blockers
   Do NOT ask minor or obvious questions. Examples of what NOT to ask:
   - For simple tasks (e.g. "add 2+5", "calculate X"): do not ask output format (integer/float), precision, or trivial preferences; use a reasonable default.
   - Do not ask about details that can be inferred from context or that do not change the deliverable.
   - Prefer assuming sensible defaults over asking the user to confirm every small detail.
   Works for both sequential (pipeline) and A2A (peer collaboration) workflows.

3. PROBLEM SOLVING: After receiving answers:
   - Synthesize all information (documents + answers)
   - Identify the core problem or need
   - Provide actionable solutions and recommendations
   - Suggest optimal approaches or workflows
   - Highlight potential challenges and mitigation strategies

4. SOLUTION-ORIENTED: Once you understand the problem:
   - Provide clear, actionable recommendations
   - Suggest specific AI agents or workflows that could solve the problem
   - Outline implementation steps
   - Identify key success factors

5. WORKFLOW MODE HINT: When you understand the workflow needed, suggest how agents should collaborate:
   - "sequential": Use when the job is a pipeline — Agent 1's output is the input to Agent 2 (step-by-step handoff). Standard agents (A2A off) are fine.
   - "async_a2a": Use when the job needs agents to work asynchronously and communicate with each other as peers (not just one output feeding the next). Recommend A2A-enabled agents.

IMPORTANT: You must respond with valid JSON only. Format your response as JSON with this exact structure:
{
    "analysis": "Comprehensive analysis of requirements, extracted data, and understanding of the problem",
    "questions": ["Question 1", "Question 2", "Question 3"] (empty array if no questions needed),
    "recommendations": ["Actionable recommendation 1", "Recommendation 2"],
    "solutions": ["Proposed solution 1", "Solution 2"] (optional, provide when problem is understood),
    "next_steps": ["Step 1", "Step 2"] (optional, provide when ready to proceed),
    "workflow_collaboration_hint": "sequential" or "async_a2a" (optional; suggest based on whether work is pipeline vs peer collaboration),
    "workflow_collaboration_reason": "One sentence explaining why." (optional; only when workflow_collaboration_hint is set)
}

When you have enough information to understand the problem, provide solutions and recommendations instead of asking more questions."""

        job_desc_part = f'Job Description: {job_description}\n\n' if job_description else ''
        conv_part = ''
        if conversation_history:
            formatted_conv = self._format_conversation(conversation_history)
            conv_part = f'\n\nPrevious Conversation History:\n{formatted_conv}\n\n'

        user_prompt = f"""Job Title: {job_title}
{job_desc_part}=== UPLOADED DOCUMENTS WITH REQUIREMENTS ===
{all_content}
{conv_part}
TASK: 
1. Perform a DEEP and COMPREHENSIVE analysis of all documents above
2. Extract ALL requirements, data structures, workflows, and business objectives
3. Identify what the user needs to accomplish
4. Ask ONLY critical questions—missing information that would change scope or outcome. Do NOT ask minor questions (e.g. for "add 2+5" do not ask integer vs float; assume a reasonable default).
5. Once you understand the problem, provide SOLUTIONS and RECOMMENDATIONS

Be thorough and solution-oriented. Prefer sensible defaults over asking trivial questions."""

        messages: List[Dict[str, Any]] = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        def _str_strip(val: Any) -> str:
            if val is None:
                return ""
            if isinstance(val, str):
                return val.strip()
            return str(val).strip()

        if conversation_history:
            for item in conversation_history:
                if item.get("type") == "question":
                    question = _str_strip(item.get("question"))
                    answer = _str_strip(item.get("answer"))
                    if question and answer:
                        messages.append({"role": "user", "content": question})
                        messages.append({"role": "assistant", "content": answer})
                elif item.get("type") == "analysis":
                    content = _str_strip(item.get("content"))
                    if content:
                        messages.append({"role": "assistant", "content": f"Analysis: {content}"})

        max_content_length = 50000
        for msg in messages:
            if len(msg.get("content", "")) > max_content_length:
                msg["content"] = msg["content"][:max_content_length] + "\n\n[Content truncated due to length...]"

        temp = getattr(settings, "AGENT_PLANNER_TEMPERATURE", None)
        if temp is None:
            temp = 0.7
        try:
            assistant_message = await planner_chat_completion(
                messages,
                temperature=float(temp),
                max_tokens=2000,
            )
        except Exception as e:
            logger.exception("Planner BRD analysis failed")
            raise Exception(f"Agent planner analysis failed: {str(e)}") from e

        try:
            parsed = json.loads(assistant_message)
            hint = parsed.get("workflow_collaboration_hint")
            if hint not in ("sequential", "async_a2a"):
                hint = None
            return {
                "analysis": parsed.get("analysis", ""),
                "questions": parsed.get("questions", []),
                "recommendations": parsed.get("recommendations", []),
                "solutions": parsed.get("solutions", []),
                "next_steps": parsed.get("next_steps", []),
                "workflow_collaboration_hint": hint,
                "workflow_collaboration_reason": (parsed.get("workflow_collaboration_reason") or "").strip() or None if hint else None,
                "raw_response": assistant_message,
            }
        except json.JSONDecodeError:
            return {
                "analysis": assistant_message,
                "questions": self._extract_questions(assistant_message),
                "recommendations": self._extract_recommendations(assistant_message),
                "solutions": [],
                "next_steps": [],
                "workflow_collaboration_hint": None,
                "workflow_collaboration_reason": None,
                "raw_response": assistant_message,
            }
    
    async def analyze_documents_and_generate_questions(
        self, 
        documents: List[Dict[str, str]], 
        job_title: str,
        job_description: Optional[str] = None,
        conversation_history: Optional[List[Dict[str, str]]] = None,
        *,
        agent_api_url: Optional[str] = None,
        agent_api_key: Optional[str] = None,
        agent_llm_model: Optional[str] = None,
        agent_temperature: Optional[float] = None,
        use_a2a: bool = False,
    ) -> Dict[str, Any]:
        """
        Extract document text and optionally analyze via hired agent.
        - When agent_api_url is provided: use hired agent for analysis and Q&A.
        - When not provided: only extract text; return extracted data and empty questions.
        """
        conversation_history = conversation_history or []
        job_title = str(job_title).strip() if job_title is not None else ""
        job_description = str(job_description).strip() if job_description is not None else None
        if job_description is not None and not job_description:
            job_description = None
        # Extract data from all documents (no inference)
        document_contents = []
        for doc in documents:
            name = doc.get("name", "Unknown")
            try:
                if doc.get("path"):
                    content = await self.read_document(doc["path"])
                else:
                    content = await self.read_file_info(doc)
            except Exception as e:
                content = f"[Error reading {name}: {str(e)}]"
            document_contents.append(f"=== {name} ===\n{content}\n")
        
        if not document_contents:
            return {
                "analysis": f"No documents could be read for job '{job_title}'.",
                "questions": [],
                "recommendations": [],
                "solutions": [],
                "next_steps": [],
                "raw_response": "",
            }
        all_content = "\n".join(document_contents)

        # Platform Agent Planner (Issue #62): admin-configured model for BRD analysis
        if is_agent_planner_configured():
            return await self._analyze_with_platform_planner(
                all_content=all_content,
                job_title=job_title,
                job_description=job_description,
                conversation_history=conversation_history,
            )
        
        # No hired agent endpoint → extraction only
        if not (agent_api_url and (agent_api_url or "").strip()):
            return {
                "analysis": f"Document text extracted for job '{job_title}'. Select and assign agents to this job to enable AI-powered analysis and Q&A.",
                "questions": [],
                "recommendations": [],
                "solutions": [],
                "next_steps": [],
                "raw_response": "",
            }
        
        # A2A path: call hired agent via A2A protocol (same semantics, different transport)
        if use_a2a:
            return await self._analyze_documents_via_a2a(
                all_content=all_content,
                job_title=job_title,
                job_description=job_description,
                conversation_history=conversation_history,
                agent_api_url=agent_api_url.strip(),
                agent_api_key=agent_api_key,
            )
        
        # OpenAI-compatible via platform adapter: route through A2A so architecture is A2A everywhere
        adapter_url = (getattr(settings, "A2A_ADAPTER_URL", None) or "").strip()
        if adapter_url:
            model = (agent_llm_model or "").strip() or "gpt-4o-mini"
            return await self._analyze_documents_via_a2a(
                all_content=all_content,
                job_title=job_title,
                job_description=job_description,
                conversation_history=conversation_history,
                agent_api_url=agent_api_url.strip(),
                agent_api_key=agent_api_key,
                adapter_url=adapter_url,
                adapter_metadata={
                    "openai_url": agent_api_url.strip(),
                    "openai_api_key": (agent_api_key or "").strip() or "",
                    "openai_model": model,
                },
            )
        
        # Hired agent: build the prompt and call agent endpoint (OpenAI-compatible, no adapter)
        system_prompt = """You are an expert AI assistant specialized in deeply understanding business requirements from documents and providing intelligent solutions.

Your primary tasks:
1. DEEP ANALYSIS: Thoroughly analyze all uploaded documents to extract:
   - Business objectives and goals
   - Technical requirements and specifications
   - Data structures, formats, and sources
   - Workflow processes and dependencies
   - Constraints, limitations, and edge cases
   - Success criteria and expected outcomes

2. INTELLIGENT QUESTIONING: Ask ONLY important, high-impact questions that:
   - Fill critical gaps in understanding (missing scope, constraints, or success criteria)
   - Clarify genuinely ambiguous requirements that affect the outcome
   - Identify real risks or blockers
   Do NOT ask minor or obvious questions. Examples of what NOT to ask:
   - For simple tasks (e.g. "add 2+5", "calculate X"): do not ask output format (integer/float), precision, or trivial preferences; use a reasonable default.
   - Do not ask about details that can be inferred from context or that do not change the deliverable.
   - Prefer assuming sensible defaults over asking the user to confirm every small detail.
   Works for both sequential (pipeline) and A2A (peer collaboration) workflows.

3. PROBLEM SOLVING: After receiving answers:
   - Synthesize all information (documents + answers)
   - Identify the core problem or need
   - Provide actionable solutions and recommendations
   - Suggest optimal approaches or workflows
   - Highlight potential challenges and mitigation strategies

4. SOLUTION-ORIENTED: Once you understand the problem:
   - Provide clear, actionable recommendations
   - Suggest specific AI agents or workflows that could solve the problem
   - Outline implementation steps
   - Identify key success factors

5. WORKFLOW MODE HINT: When you understand the workflow needed, suggest how agents should collaborate:
   - "sequential": Use when the job is a pipeline — Agent 1's output is the input to Agent 2 (step-by-step handoff). Standard agents (A2A off) are fine.
   - "async_a2a": Use when the job needs agents to work asynchronously and communicate with each other as peers (not just one output feeding the next). Recommend A2A-enabled agents.

IMPORTANT: You must respond with valid JSON only. Format your response as JSON with this exact structure:
{
    "analysis": "Comprehensive analysis of requirements, extracted data, and understanding of the problem",
    "questions": ["Question 1", "Question 2", "Question 3"] (empty array if no questions needed),
    "recommendations": ["Actionable recommendation 1", "Recommendation 2"],
    "solutions": ["Proposed solution 1", "Solution 2"] (optional, provide when problem is understood),
    "next_steps": ["Step 1", "Step 2"] (optional, provide when ready to proceed),
    "workflow_collaboration_hint": "sequential" or "async_a2a" (optional; suggest based on whether work is pipeline vs peer collaboration),
    "workflow_collaboration_reason": "One sentence explaining why." (optional; only when workflow_collaboration_hint is set)
}

When you have enough information to understand the problem, provide solutions and recommendations instead of asking more questions."""
        
        # Build user prompt
        job_desc_part = f'Job Description: {job_description}\n\n' if job_description else ''
        conv_part = ''
        if conversation_history:
            formatted_conv = self._format_conversation(conversation_history)
            conv_part = f'\n\nPrevious Conversation History:\n{formatted_conv}\n\n'
        
        user_prompt = f"""Job Title: {job_title}
{job_desc_part}=== UPLOADED DOCUMENTS WITH REQUIREMENTS ===
{all_content}
{conv_part}
TASK: 
1. Perform a DEEP and COMPREHENSIVE analysis of all documents above
2. Extract ALL requirements, data structures, workflows, and business objectives
3. Identify what the user needs to accomplish
4. Ask ONLY critical questions—missing information that would change scope or outcome. Do NOT ask minor questions (e.g. for "add 2+5" do not ask integer vs float; assume a reasonable default).
5. Once you understand the problem, provide SOLUTIONS and RECOMMENDATIONS

Be thorough and solution-oriented. Prefer sensible defaults over asking trivial questions."""
        
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        # Add conversation history if provided - only include Q&A pairs with both question and answer
        def _str_strip(val: Any) -> str:
            if val is None:
                return ""
            if isinstance(val, str):
                return val.strip()
            return str(val).strip()

        if conversation_history:
            for item in conversation_history:
                if item.get("type") == "question":
                    question = _str_strip(item.get("question"))
                    answer = _str_strip(item.get("answer"))
                    if question and answer:
                        messages.append({"role": "user", "content": question})
                        messages.append({"role": "assistant", "content": answer})
                elif item.get("type") == "analysis":
                    content = _str_strip(item.get("content"))
                    if content:
                        messages.append({"role": "assistant", "content": f"Analysis: {content}"})
        
        try:
            max_content_length = 50000
            for msg in messages:
                if len(msg.get("content", "")) > max_content_length:
                    msg["content"] = msg["content"][:max_content_length] + "\n\n[Content truncated due to length...]"
            
            model = (agent_llm_model or "").strip() or "gpt-4o-mini"
            temperature = agent_temperature if agent_temperature is not None else 0.7
            payload = {
                "model": model,
                "messages": messages,
                "temperature": temperature,
                "max_tokens": 2000,
            }
            headers = {"Content-Type": "application/json"}
            if agent_api_key and (agent_api_key or "").strip():
                headers["Authorization"] = f"Bearer {(agent_api_key or '').strip()}"
            
            async with httpx.AsyncClient(timeout=120.0, verify=False) as client:
                response = await client.post(
                    agent_api_url.strip(),
                    json=payload,
                    headers=headers,
                )
                response.raise_for_status()
                result = response.json()
                raw_content = result["choices"][0]["message"].get("content")
                # OpenAI can return content as string or list of parts (e.g. multimodal)
                if isinstance(raw_content, list):
                    assistant_message = " ".join(
                        p.get("text", p.get("content", "")) if isinstance(p, dict) else str(p)
                        for p in raw_content
                    )
                elif isinstance(raw_content, str):
                    assistant_message = raw_content
                else:
                    assistant_message = str(raw_content or "")
                
                try:
                    parsed = json.loads(assistant_message)
                    hint = parsed.get("workflow_collaboration_hint")
                    if hint not in ("sequential", "async_a2a"):
                        hint = None
                    return {
                        "analysis": parsed.get("analysis", ""),
                        "questions": parsed.get("questions", []),
                        "recommendations": parsed.get("recommendations", []),
                        "solutions": parsed.get("solutions", []),
                        "next_steps": parsed.get("next_steps", []),
                        "workflow_collaboration_hint": hint,
                        "workflow_collaboration_reason": (parsed.get("workflow_collaboration_reason") or "").strip() or None if hint else None,
                        "raw_response": assistant_message,
                    }
                except json.JSONDecodeError:
                    return {
                        "analysis": assistant_message,
                        "questions": self._extract_questions(assistant_message),
                        "recommendations": self._extract_recommendations(assistant_message),
                        "solutions": [],
                        "next_steps": [],
                        "workflow_collaboration_hint": None,
                        "workflow_collaboration_reason": None,
                        "raw_response": assistant_message,
                    }
        except httpx.HTTPStatusError as e:
            error_detail = f"Hired agent API error: {e.response.status_code}"
            try:
                error_detail += f" - {e.response.text[:500]}"
            except Exception:
                pass
            raise Exception(error_detail)
        except httpx.RequestError as e:
            raise Exception(f"Hired agent API request error: {str(e)}")
        except Exception as e:
            raise Exception(f"Error analyzing documents: {str(e)}")
    
    async def process_user_response(
        self,
        user_answer: str,
        documents: List[Dict[str, str]],
        job_title: str,
        job_description: Optional[str],
        conversation_history: List[Dict[str, str]],
        *,
        agent_api_url: Optional[str] = None,
        agent_api_key: Optional[str] = None,
        agent_llm_model: Optional[str] = None,
        agent_temperature: Optional[float] = None,
        use_a2a: bool = False,
    ) -> Dict[str, Any]:
        """Process user's answer and generate follow-up questions or final recommendations via hired agent."""
        last_question = None
        for item in reversed(conversation_history):
            if item.get("type") == "question" and not item.get("answer"):
                last_question = item.get("question") or ""
                break
        
        updated_history = []
        answer_added = False
        for item in conversation_history:
            if item.get("type") == "question" and not item.get("answer") and not answer_added:
                updated_item = item.copy()
                updated_item["answer"] = user_answer
                updated_history.append(updated_item)
                answer_added = True
            else:
                updated_history.append(item)
        
        if not answer_added and last_question:
            updated_history.append({
                "type": "question",
                "question": last_question,
                "answer": user_answer
            })
        
        return await self.analyze_documents_and_generate_questions(
            documents,
            job_title,
            job_description,
            updated_history,
            agent_api_url=agent_api_url,
            agent_api_key=agent_api_key,
            agent_llm_model=agent_llm_model,
            agent_temperature=agent_temperature,
            use_a2a=use_a2a,
        )
    
    async def generate_workflow_clarification_questions(
        self,
        job_title: str,
        job_description: Optional[str],
        documents_content: List[Dict[str, Any]],
        workflow_tasks: List[Dict[str, Any]],
        conversation_history: List[Dict[str, Any]],
        *,
        agent_api_url: Optional[str] = None,
        agent_api_key: Optional[str] = None,
        agent_llm_model: Optional[str] = None,
        agent_temperature: Optional[float] = None,
        use_a2a: bool = False,
        only_step: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Generate clarifying questions for the end user based on the workflow (assigned tasks),
        BRD documents, and job prompt. Used in the Q&A step after Build Workflow so agents
        can get requirements clarified before execution.
        When only_step is set, scope is that workflow step only (assigned agent's endpoint).
        Returns {"questions": ["...", ...]}.
        """
        max_q = 2 if only_step is not None else 3
        if only_step is not None:
            tasks_for_prompt = [
                {
                    "step_order": only_step.get("step_order", 1),
                    "agent_name": only_step.get("agent_name", "Agent"),
                    "assigned_task": only_step.get("assigned_task", ""),
                }
            ]
        else:
            if not workflow_tasks:
                return {"questions": []}
            tasks_for_prompt = workflow_tasks

        job_title = (job_title or "").strip()
        job_description = (job_description or "").strip() or ""
        docs_text = ""
        if documents_content:
            docs_text = "\n\n".join(
                f"=== {d.get('name', 'Unknown')} ===\n{(d.get('content') or '')[:8000]}"
                for d in documents_content
            )
        workflow_text = "\n".join(
            f"Step {t.get('step_order', i+1)} ({t.get('agent_name', 'Agent')}): {t.get('assigned_task', '')}"
            for i, t in enumerate(tasks_for_prompt)
        )
        conv_text = self._format_conversation(conversation_history) if conversation_history else "(none)"
        user_prompt = f"""JOB TITLE: {job_title}
JOB DESCRIPTION: {job_description}

BRD / REQUIREMENT DOCUMENTS:
{docs_text or '(no documents)'}

WORKFLOW (assigned task per agent):
{workflow_text}

EXISTING Q&A:
{conv_text}

TASK: Generate clarification questions ONLY (no answers, no analysis summary, no solutions).
Your goal is to help each assigned agent execute its own task correctly by asking targeted requirement questions.
{"You are generating questions ONLY for the single workflow step listed above (you represent that assigned agent)." if only_step is not None else ""}

RULES:
1) Questions must be tied to specific workflow steps/agents (based on assigned tasks above).
2) Ask ONLY critical execution questions:
   - missing required input that cannot be inferred,
   - real ambiguity/conflict that changes execution/result,
   - missing hard constraint (security/compliance/SLA/limits/format that is mandatory).
3) Do NOT provide answers, computed results, recommendations, or implementation steps in this stage.
4) If and only if all assigned tasks are executable with clear requirements, return an empty array.
5) Prefer concise, actionable questions. Avoid generic or conversational filler.
6) Maximum {max_q} questions total. If more than {max_q} are possible, keep only the highest-impact blockers.
7) Before returning each question, apply this quality gate:
   - Would the answer materially change implementation/scope for a workflow step?
   - Is this information truly unavailable from BRD + job + existing Q&A?
   - Is this blocking execution now?
   If any answer is "no", DO NOT ask that question.

DO NOT ask:
- output format (integer/float) unless explicitly required by a downstream step,
- display/wording preferences,
- \"which method/tool do you prefer?\",
- context/background questions that do not affect execution,
- additional nice-to-have requests unrelated to assigned steps.
- questions already answered in EXISTING Q&A.

GOOD (critical) examples:
- "[Step 2 - Pricing Agent] What tax jurisdiction should be applied for VAT calculation?"
- "[Step 1 - Data Agent] Which source-of-truth table should be used when CRM and ERP values conflict?"

BAD (silly/non-critical) examples:
- "Do you want the output as integer or float?"
- "Any preferred method/tool?"
- "Can you provide context for this result?"

Return ONLY valid JSON with this exact shape:
{{"questions": ["[Step 1 - Agent Name] ...?", "[Step 2 - Agent Name] ...?"]}}
or
{{"questions": []}}
No markdown. No extra keys. No explanation text."""

        if not (agent_api_url and (agent_api_url or "").strip()):
            return {"questions": []}

        adapter_url = (getattr(settings, "A2A_ADAPTER_URL", None) or "").strip()
        if use_a2a and not adapter_url:
            # Direct A2A to agent
            input_data = {"job_title": job_title, "job_description": job_description, "prompt": user_prompt}
            result = await execute_via_a2a(
                agent_api_url.strip(),
                input_data,
                api_key=(agent_api_key or "").strip() or None,
                blocking=True,
                timeout=120.0,
            )
            content = (result.get("content") or "").strip()
        elif adapter_url:
            model = (agent_llm_model or "").strip() or "gpt-4o-mini"
            input_data = {"job_title": job_title, "job_description": job_description, "prompt": user_prompt}
            result = await execute_via_a2a(
                adapter_url,
                input_data,
                api_key=None,
                blocking=True,
                timeout=120.0,
                adapter_metadata={
                    "openai_url": agent_api_url.strip(),
                    "openai_api_key": (agent_api_key or "").strip() or "",
                    "openai_model": model,
                },
            )
            content = (result.get("content") or "").strip()
        else:
            # Direct OpenAI
            model = (agent_llm_model or "").strip() or "gpt-4o-mini"
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": "You are in CLARIFICATION mode for workflow execution. Generate ONLY execution-critical questions mapped to workflow steps/agents. NEVER provide answers, calculations, solutions, recommendations, summaries, or next steps. Ask only when a required input/constraint/ambiguity blocks correct execution for a step. Enforce strict quality gate: each question must materially change implementation/scope and be currently blocking and not already answered. MAX 3 questions. Reject trivial questions (format preference, method preference, generic context). If all steps are executable, return empty array. Return only valid JSON: {\"questions\": [\"[Step N - Agent Name] ...?\"]} or {\"questions\": []}. No markdown. No extra keys."},
                    {"role": "user", "content": user_prompt[:50000]},
                ],
                "temperature": agent_temperature if agent_temperature is not None else 0.5,
                "max_tokens": 1500,
            }
            headers = {"Content-Type": "application/json"}
            if agent_api_key and (agent_api_key or "").strip():
                headers["Authorization"] = f"Bearer {(agent_api_key or '').strip()}"
            async with httpx.AsyncClient(timeout=120.0, verify=httpx_verify_parameter()) as client:
                resp = await client.post(agent_api_url.strip(), json=payload, headers=headers)
                resp.raise_for_status()
                data = resp.json()
                raw = (data.get("choices") or [{}])[0].get("message") or {}
                content = (raw.get("content") or "").strip()
                if isinstance(content, list):
                    content = " ".join(
                        p.get("text", p.get("content", "")) if isinstance(p, dict) else str(p)
                        for p in content
                    )

        if not content:
            return {"questions": []}
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
            content = content.strip()
        try:
            parsed = json.loads(content)
            questions = parsed.get("questions")
            if isinstance(questions, list):
                return {
                    "questions": self._filter_critical_questions(
                        [str(q).strip() for q in questions if str(q).strip()],
                        max_count=max_q,
                    )
                }
        except (json.JSONDecodeError, TypeError):
            pass
        return {
            "questions": self._filter_critical_questions(
                self._extract_questions(content)[:max_q],
                max_count=max_q,
            )
        }

    def _format_conversation(self, history: List[Dict[str, Any]]) -> str:
        """Format conversation history for the prompt"""
        formatted = []
        for i, item in enumerate(history, 1):
            q = item.get("question")
            a = item.get("answer")
            question = q.strip() if isinstance(q, str) else (str(q) if q else "")
            answer = a.strip() if isinstance(a, str) else (str(a) if a else "")
            if question or answer:
                if question:
                    formatted.append(f"Q{i}: {question}")
                if answer:
                    formatted.append(f"A{i}: {answer}")
        return "\n".join(formatted)
    
    def _extract_questions(self, text: str) -> List[str]:
        """Extract questions from text if JSON parsing fails"""
        import re
        # Look for lines ending with ?
        questions = re.findall(r'[^.!?]*\?', text)
        return [q.strip() for q in questions if q.strip()][:3]  # Limit to 3 questions
    
    def _extract_recommendations(self, text: str) -> List[str]:
        """Extract recommendations/solutions from text if JSON parsing fails"""
        import re
        recommendations = []
        # Look for bullet points, numbered lists, or lines starting with keywords
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            # Match bullet points, numbered items, or lines with recommendation keywords
            if (line.startswith('-') or line.startswith('*') or 
                re.match(r'^\d+[\.\)]', line) or
                any(keyword in line.lower() for keyword in ['recommend', 'suggest', 'solution', 'should', 'consider'])):
                # Clean up the line
                cleaned = re.sub(r'^[-*\d+\.\)]\s*', '', line)
                if cleaned and len(cleaned) > 10:  # Only include substantial recommendations
                    recommendations.append(cleaned)
        return recommendations[:5]  # Limit to 5 recommendations

    def _filter_critical_questions(self, questions: List[str], max_count: int = 3) -> List[str]:
        """
        Keep only high-signal clarification questions.
        Filters out common trivial/silly prompt artifacts and limits to top max_count.
        """
        if not questions:
            return []

        blocked_phrases = (
            "integer or float",
            "int or float",
            "output format",
            "display format",
            "preferred method",
            "preferred tool",
            "any specific method",
            "context for this result",
            "additional operations",
            "precision level",
            "anything else",
        )
        keep: List[str] = []
        seen = set()
        for q in questions:
            clean = (q or "").strip()
            if not clean:
                continue
            low = clean.lower()
            if any(p in low for p in blocked_phrases):
                continue
            if low in seen:
                continue
            seen.add(low)
            keep.append(clean)
            if len(keep) >= max_count:
                break
        return keep