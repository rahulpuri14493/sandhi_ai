"""
Integration coverage for full document-upload steps in S3 mode.

This test validates the entire flow:
1) authenticated business user uploads a BRD (.docx)
2) API response metadata is redacted
3) DB stores internal S3 metadata (bucket/key/storage)
4) object is present and readable in MinIO/S3
5) file can be downloaded back via API
6) uploaded object is not persisted as a local uploads/jobs file
"""

import io
import json
import os
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from models.job import Job


def _auth_headers(user) -> dict:
    return {"Authorization": f"Bearer {user['token']}"}


def _build_brd_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("BRD - Upload Steps Test", level=1)
    doc.add_paragraph("Business Goal: verify S3/MinIO-backed upload flow.")
    doc.add_paragraph("Requirement 1: upload BRD document.")
    doc.add_paragraph("Requirement 2: persist in S3-compatible object storage.")
    doc.add_paragraph("Requirement 3: retrieve document through API.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _s3_enabled() -> bool:
    backend = (os.environ.get("OBJECT_STORAGE_BACKEND") or "local").strip().lower()
    return backend == "s3"


@pytest.mark.skipif(not _s3_enabled(), reason="Requires OBJECT_STORAGE_BACKEND=s3")
def test_document_upload_all_steps(
    integration_client: TestClient,
    business_user,
    integration_db_session,
):
    # ---- Step 1: upload BRD as authenticated business user ----
    brd_name = "upload-steps-brd.docx"
    brd_bytes = _build_brd_docx_bytes()
    create_resp = integration_client.post(
        "/api/jobs",
        data={"title": "BRD Upload Steps", "description": "End-to-end upload validation"},
        files={
            "files": (
                brd_name,
                brd_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers=_auth_headers(business_user),
    )
    assert create_resp.status_code == 201, create_resp.text
    body = create_resp.json()
    assert body["id"] > 0
    assert body["files"] and len(body["files"]) == 1

    public_file = body["files"][0]
    job_id = body["id"]
    file_id = public_file["id"]
    file_name = public_file["name"]
    assert file_name.endswith(".docx")

    # ---- Step 2: verify API response is redacted (no private storage fields) ----
    for private_key in ("path", "bucket", "key", "storage"):
        assert private_key not in public_file

    # ---- Step 3: verify DB stores internal S3 metadata ----
    job = integration_db_session.query(Job).filter(Job.id == job_id).first()
    assert job is not None
    stored_files = json.loads(job.files or "[]")
    assert len(stored_files) == 1
    internal_file = stored_files[0]
    assert internal_file["id"] == file_id
    assert internal_file["name"] == file_name
    assert internal_file["storage"] == "s3"
    assert internal_file["bucket"] == os.environ.get("S3_BUCKET")
    assert internal_file["key"].startswith(f"jobs/{job_id}/")
    assert internal_file["key"].endswith(f"_{file_name}")
    assert "path" not in internal_file

    # ---- Step 4: verify object exists in MinIO/S3 and content matches ----
    import boto3
    from botocore.config import Config

    s3_client = boto3.client(
        "s3",
        endpoint_url=os.environ.get("S3_ENDPOINT_URL"),
        aws_access_key_id=os.environ.get("S3_ACCESS_KEY_ID"),
        aws_secret_access_key=os.environ.get("S3_SECRET_ACCESS_KEY"),
        region_name=os.environ.get("S3_REGION", "us-east-1"),
        config=Config(
            signature_version=(os.environ.get("S3_SIGNATURE_VERSION") or "s3v4").strip(),
            s3={"addressing_style": (os.environ.get("S3_ADDRESSING_STYLE") or "path").strip()},
        ),
    )
    obj = s3_client.get_object(Bucket=internal_file["bucket"], Key=internal_file["key"])
    s3_bytes = obj["Body"].read()
    assert s3_bytes == brd_bytes

    # ---- Step 5: verify API download returns the same content ----
    dl_resp = integration_client.get(
        f"/api/jobs/{job_id}/files/{file_id}",
        headers=_auth_headers(business_user),
    )
    assert dl_resp.status_code == 200, dl_resp.text
    assert dl_resp.content == brd_bytes
    assert (
        dl_resp.headers.get("content-type", "")
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # ---- Step 6: verify no local uploads/jobs object persisted for this file ----
    expected_local = Path("uploads/jobs") / f"{file_id}_{file_name}"
    assert not expected_local.exists(), f"Unexpected local file persisted: {expected_local}"
